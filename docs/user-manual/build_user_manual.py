from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
ASSET_DIR = BASE_DIR / "assets"
OUTPUT_PATH = BASE_DIR / "Hyper用户使用手册_小程序与网页端_20260810_正式版.docx"

NAVY = "0B1F33"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
GOLD = "D7A84A"
PINK = "FF1654"
INK = "17212B"
MUTED = "66717E"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF7E6"
PALE_RED = "FFF0F3"
PALE_GREEN = "EDF8F2"
LINE = "C6D3E1"
WHITE = "FFFFFF"
BLACK = "000000"

# Use one full-Unicode family for all script slots. LibreOffice's headless
# renderer can ignore w:eastAsia when w:ascii points at Calibri, which turns
# Chinese glyphs into tofu even though macOS has PingFang installed.
FONT_LATIN = "Arial Unicode MS"
FONT_EAST_ASIA = "Arial Unicode MS"

figure_counter = 0
step_counter = 0


def set_run_font(run, *, name: str = FONT_LATIN, east_asia: str = FONT_EAST_ASIA,
                 size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    r_fonts = run._element.get_or_add_rPr().rFonts
    for slot, family in (("ascii", name), ("hAnsi", name),
                         ("eastAsia", east_asia), ("cs", name)):
        r_fonts.set(qn(f"w:{slot}"), family)
    r_fonts.set(qn("w:hint"), "eastAsia")
    for theme_slot in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{theme_slot}")
        if key in r_fonts.attrib:
            del r_fonts.attrib[key]
    r_pr = run._element.get_or_add_rPr()
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start),
                                      ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, top: tuple[str, int] | None = None,
                    bottom: tuple[str, int] | None = None,
                    start: tuple[str, int] | None = None,
                    end: tuple[str, int] | None = None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in (("top", top), ("bottom", bottom),
                       ("start", start), ("end", end)):
        if spec is None:
            continue
        color, size = spec
        tag = f"w:{edge}"
        edge_el = tc_borders.find(qn(tag))
        if edge_el is None:
            edge_el = OxmlElement(tag)
            tc_borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), str(size))
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "nil")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_no_split(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_together = value


def add_field(paragraph, instruction: str, fallback: str = "1") -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_paragraph(doc: Document, text: str = "", *, style: str | None = None,
                  bold_prefix: str | None = None) -> object:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=INK)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> object:
    global step_counter
    step_counter = 0
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_run_font(r, bold=True, color=BLUE if level < 3 else BLUE_DARK)
    set_keep_with_next(p)
    return p


def add_step(doc: Document, text: str) -> object:
    global step_counter
    step_counter += 1
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    number = p.add_run(f"{step_counter}. ")
    set_run_font(number, bold=True, color=BLUE_DARK)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_bullet(doc: Document, text: str, level: int = 0) -> object:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_callout(doc: Document, title: str, body: str, *, kind: str = "info",
                compact: bool = False, trailing_space: bool = True) -> None:
    palette = {
        "info": (PALE_BLUE, BLUE),
        "warn": (PALE_GOLD, GOLD),
        "danger": (PALE_RED, PINK),
        "success": (PALE_GREEN, "2D8B57"),
    }
    fill, accent = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    remove_table_borders(table)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(
        cell,
        top=80 if compact else 120,
        start=150 if compact else 170,
        bottom=80 if compact else 120,
        end=150 if compact else 170,
    )
    set_cell_border(cell, start=(accent, 18))
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10 if compact else 10.5, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12 if compact else 1.18
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.2 if compact else 9.5, color=INK)
    if trailing_space:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_data_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
                   widths: Sequence[float] | None = None,
                   trailing_space: bool = True) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    header = table.rows[0]
    set_repeat_table_header(header)
    header._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    for idx, (text, width) in enumerate(zip(headers, widths)):
        cell = header.cells[idx]
        cell.width = Inches(width)
        set_cell_shading(cell, PALE_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=BLUE_DARK)
    for row_data in rows:
        row = table.add_row()
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        cells = row.cells
        for idx, (text, width) in enumerate(zip(row_data, widths)):
            cell = cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(text))
            set_run_font(r, size=9.2, color=INK)
    table_pr = table._tbl.tblPr
    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    if trailing_space:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_picture_to_paragraph(paragraph, path: Path, width: float, alt: str):
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("title", alt)
    doc_pr.set("descr", alt)
    return inline


def add_figure(doc: Document, filename: str, caption: str, *, width: float = 6.25,
               alt: str | None = None) -> int:
    global figure_counter
    figure_counter += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    set_no_split(p)
    set_keep_with_next(p)
    add_picture_to_paragraph(p, ASSET_DIR / filename, width,
                             alt or caption)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(7)
    set_no_split(cap)
    r = cap.add_run(f"图 {figure_counter}  {caption}")
    set_run_font(r, size=8.5, color=MUTED)
    return figure_counter


def add_figure_pair(doc: Document, figures: Sequence[tuple[str, str, str]]) -> None:
    global figure_counter
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    for idx, (filename, caption, alt) in enumerate(figures):
        figure_counter += 1
        cell = table.cell(0, idx)
        cell.width = Inches(3.2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(cell, top=0, start=40, bottom=0, end=40)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        add_picture_to_paragraph(p, ASSET_DIR / filename, 2.4, alt)
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(0)
        r = cap.add_run(f"图 {figure_counter}  {caption}")
        set_run_font(r, size=8.3, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text.upper())
    set_run_font(r, size=8.5, bold=True, color=PINK)
    r.font.letter_spacing = Pt(1.2) if hasattr(r.font, "letter_spacing") else None


def add_part_title(doc: Document, number: str, title: str, subtitle: str) -> None:
    global step_counter
    step_counter = 0
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(number)
    set_run_font(r, size=11, bold=True, color=PINK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(title)
    set_run_font(r2, size=30, bold=True, color=NAVY)
    rule = doc.add_table(rows=1, cols=2)
    rule.alignment = WD_TABLE_ALIGNMENT.LEFT
    rule.autofit = False
    remove_table_borders(rule)
    rule.cell(0, 0).width = Inches(1.1)
    rule.cell(0, 1).width = Inches(5.4)
    set_cell_shading(rule.cell(0, 0), PINK)
    set_cell_shading(rule.cell(0, 1), PALE_BLUE)
    for cell in rule.rows[0].cells:
        set_cell_margins(cell, top=25, start=0, bottom=25, end=0)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(24)
    p3.paragraph_format.space_after = Pt(8)
    r3 = p3.add_run(subtitle)
    set_run_font(r3, size=13, color=BLUE_DARK)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = FONT_LATIN
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = FONT_LATIN
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = FONT_LATIN
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(BLUE_DARK)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    bullet2 = styles["List Bullet 2"]
    bullet2.font.name = FONT_LATIN
    bullet2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    bullet2.font.size = Pt(10.5)
    bullet2.paragraph_format.left_indent = Inches(0.625)
    bullet2.paragraph_format.first_line_indent = Inches(-0.188)
    bullet2.paragraph_format.space_after = Pt(3)
    bullet2.paragraph_format.line_spacing = 1.2

    caption = styles["Caption"]
    caption.font.name = FONT_LATIN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("HYPER  /  用户使用手册")
    set_run_font(hr, size=8.5, bold=True, color=MUTED)
    hp_pr = hp._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    borders.append(bottom)
    hp_pr.append(borders)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("2026年8月版  ·  ")
    set_run_font(fr, size=8.5, color=MUTED)
    add_field(fp, "PAGE", "1")

    settings = doc.settings._element
    theme_font_lang = settings.find(qn("w:themeFontLang"))
    if theme_font_lang is None:
        theme_font_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_font_lang)
    theme_font_lang.set(qn("w:val"), "zh-CN")
    theme_font_lang.set(qn("w:eastAsia"), "zh-CN")
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    props = doc.core_properties
    props.title = "Hyper 用户使用手册：小程序与网页端"
    props.subject = "小程序用户端、主办方端、网页商家端与平台管理端操作说明"
    props.author = "Hyper 项目组"
    props.keywords = "Hyper, 用户手册, 小程序, 商家后台, 管理后台"
    props.comments = "Hyper 用户使用手册，2026年8月版"


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("HYPER")
    set_run_font(r, size=12, bold=True, color=NAVY)
    r = p.add_run("  /  用户使用手册")
    set_run_font(r, size=8.5, bold=True, color=MUTED)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), GOLD)
    borders.append(bottom)
    p_pr.append(borders)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(62)
    spacer.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("USER GUIDE  /  2026.08")
    set_run_font(r, size=10, bold=True, color=PINK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Hyper 用户使用手册")
    set_run_font(r, size=29, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("小程序 · 网页商家端 · 平台管理端")
    set_run_font(r, size=16, color=BLUE_DARK)

    accent = doc.add_table(rows=1, cols=2)
    accent.alignment = WD_TABLE_ALIGNMENT.LEFT
    accent.autofit = False
    remove_table_borders(accent)
    accent.cell(0, 0).width = Inches(1.1)
    accent.cell(0, 1).width = Inches(5.4)
    set_cell_shading(accent.cell(0, 0), GOLD)
    set_cell_shading(accent.cell(0, 1), PALE_BLUE)
    for c in accent.rows[0].cells:
        set_cell_margins(c, top=28, start=0, bottom=28, end=0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("覆盖活动浏览、报名购票、票券核销、主办方运营、商家管理与平台治理的操作指南。")
    set_run_font(r, size=12, color=INK)

    chips = doc.add_table(rows=1, cols=3)
    chips.alignment = WD_TABLE_ALIGNMENT.LEFT
    chips.autofit = False
    remove_table_borders(chips)
    for idx, label in enumerate(("小程序用户", "主办方 / 核销员", "网页商家 / 管理员")):
        c = chips.cell(0, idx)
        c.width = Inches(2.08)
        set_cell_shading(c, PALE_BLUE if idx != 1 else PALE_GOLD)
        set_cell_margins(c, top=110, start=100, bottom=110, end=100)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=9.5, bold=True, color=BLUE_DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("版本 1.0  ·  2026年8月10日")
    set_run_font(r, size=10.5, bold=True, color=NAVY)


def build_manual() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_page_break()
    add_heading(doc, "阅读说明", 1)
    add_paragraph(
        doc,
        "本手册面向活动参与者、主办方、核销员、商家运营人员和平台管理员。"
        "正文按角色拆分，可直接跳到与自己相关的章节。"
    )
    add_data_table(
        doc,
        ["对象", "入口", "本手册重点"],
        [
            ["活动参与者", "微信小程序", "活动浏览、报名购票、订单与电子票、退款、消息及积分"],
            ["主办方 / 核销员", "小程序 → 我的 → 管理后台 / 核销入口", "入驻申请、活动发布、销售分析、订单管理、票券核销及账户管理"],
            ["商家运营人员", "网页活动管理系统", "活动、订单售后、资金、运营、权限设置"],
            ["平台管理员", "网页登录页 → 管理员登录", "审核、订单与售后、核销、内容、财务积分、权限治理"],
        ],
        [1.2, 1.8, 3.5],
    )
    add_heading(doc, "适用范围", 2)
    add_data_table(
        doc,
        ["终端", "适用对象", "主要功能"],
        [
            ["微信小程序", "活动参与者、主办方、核销员", "活动浏览、报名购票、订单管理、活动发布与票券核销"],
            ["网页活动管理系统", "商家运营人员、平台管理员", "活动运营、订单售后、账户资金、审核及权限管理"],
            ["手册版本", "2026年8月版", "页面名称、字段和功能以正式发布版本为准"],
        ],
        [1.45, 2.15, 2.9],
    )
    add_callout(
        doc,
        "界面说明",
        "本手册所示界面用于说明操作路径。因账号权限、业务状态及版本更新，实际菜单、字段和提示可能有所不同，请以正式发布页面为准。",
        kind="info",
    )
    add_callout(
        doc,
        "功能范围",
        "商家等级规则由平台统一维护，商家仅可查看，不能自行编辑。当前版本暂不提供活动分销和活动抽奖功能。",
        kind="warn",
    )

    add_heading(doc, "快速目录", 1)
    add_data_table(
        doc,
        ["部分", "适用角色", "包含内容"],
        [
            ["01 小程序用户端", "活动参与者", "发现、广场、登录、下单、电子票、售后、消息、积分"],
            ["02 小程序主办方端", "主办方、核销员", "入驻、活动发布、销售订单、核销、账户资金"],
            ["03 网页商家端", "运营、财务、客服、管理员", "登录、活动、订单售后、账户、运营、权限"],
            ["04 平台管理端", "平台管理员", "数据概览、审核、订单、核销、内容、财务积分"],
            ["05 状态与常见问题", "所有角色", "状态含义、故障排查、安全建议"],
        ],
        [1.45, 1.55, 3.5],
    )
    add_heading(doc, "推荐使用顺序", 2)
    add_step(doc, "先确认自己的角色与入口；网页端菜单会随账号权限而变化。")
    add_step(doc, "首次使用先完成登录、手机号授权或商家入驻，再进入具体业务操作。")
    add_step(doc, "涉及付款、退款、核销、提现、审核的操作，提交前再次核对订单号、金额和对象。")
    add_step(doc, "操作结果以状态标签和消息中心通知为准；按钮提交成功并不代表业务流程已经完成。")

    # Part 1: mini-program user
    add_part_title(doc, "PART 01", "小程序用户端", "适用于浏览活动、报名购票、查看电子票与发起售后的活动参与者。")
    add_heading(doc, "1.1 首次打开与登录", 1)
    add_paragraph(doc, "未登录时可以浏览首页和广场；下单、关注、发布动态、查看订单等操作需要先完成登录。")
    add_step(doc, "在底部导航点击“我的”，再点击头像区域或需要登录的功能。")
    add_step(doc, "按照微信提示完成授权；如页面要求手机号，选择微信绑定手机号或按提示验证。")
    add_step(doc, "阅读并同意相关协议后完成登录。登录成功后，“我的”页会显示订单、积分、入驻和管理后台等入口。")
    add_callout(doc, "账号安全", "验证码、支付密码和核销二维码均不要转发给他人。更换手机号或设备后，如出现身份不一致，请先退出再重新登录。", kind="warn")

    add_heading(doc, "1.2 首页活动浏览", 1)
    add_paragraph(doc, "首页同时提供地图与活动列表。定位可用时，地图会显示附近派对或场地；下方列表用于快速浏览。")
    add_step(doc, "允许定位后查看附近活动；拒绝定位时仍可通过列表、搜索或筛选查找。")
    add_step(doc, "使用搜索输入活动名、场地或关键词；使用筛选按日期、类型或距离缩小范围。")
    add_step(doc, "点击活动卡片进入详情，核对时间、地点、主办方、退票说明、实名与未成年人限制。")
    add_step(doc, "选择票种和数量后提交订单；在微信支付页确认金额并完成支付。")
    add_figure_pair(
        doc,
        [
            ("mini_home.png", "首页：地图与活动列表", "小程序首页，包含地图区域和活动列表"),
            ("mini_square.png", "广场：动态瀑布流", "小程序广场页面，展示用户与主办方动态"),
        ],
    )

    add_heading(doc, "1.3 广场、关注与互动", 1)
    add_paragraph(doc, "广场用于浏览用户或主办方发布的动态。点击卡片可进入动态详情，再查看作者主页或关联活动。")
    add_bullet(doc, "关注：在作者主页点击“关注”；可在“我的”相关入口查看关注关系。")
    add_bullet(doc, "点赞与评论：登录后操作；请勿发布违法、侵权、广告或包含隐私的信息。")
    add_bullet(doc, "发布动态：点击广场发布入口，上传图片、填写文字，按需关联活动后提交。")
    add_bullet(doc, "举报或屏蔽：如页面显示相应入口，可按提示提交；处理结果以平台通知为准。")

    add_heading(doc, "1.4 个人中心、订单与电子票", 1)
    add_paragraph(doc, "进入“我的 → 订单”查看全部订单，或按状态筛选待支付、待使用、已使用、退款中等订单。")
    add_step(doc, "打开订单列表并选择目标订单；先核对活动名称、场次、票种、数量和订单号。")
    add_step(doc, "待支付订单：在有效期内继续付款；超时后以订单状态为准，不要重复支付。")
    add_step(doc, "待使用订单：进入详情查看电子票或二维码；到场时调高屏幕亮度，逐张出示。")
    add_step(doc, "需要退款时，在订单详情发起售后，选择原因并提交；审核期间不要重复申请。")
    add_figure_pair(
        doc,
        [
            ("mini_user_center.png", "“我的”页的访客状态与登录入口", "小程序我的页面，展示未登录状态和功能入口"),
            ("mini_orders.png", "订单列表与状态筛选", "小程序订单列表，包含多种订单状态标签"),
        ],
    )
    add_callout(doc, "电子票使用", "二维码通常对应具体票券。截图、转发或提前暴露二维码可能带来被他人冒用的风险；核销成功后票券不可再次使用。", kind="danger")

    add_heading(doc, "1.5 消息、积分与个人资料", 1)
    add_bullet(doc, "消息：底部进入“消息”，可查看系统通知、互动通知、支付通知及会话；红点表示未读。")
    add_bullet(doc, "积分：从“我的 → 积分”查看余额与明细；积分获得和使用以页面规则为准。")
    add_bullet(doc, "个人主页：可维护头像、昵称等资料，并查看已发布动态或关注关系。")
    add_bullet(doc, "订阅提醒：当页面请求微信订阅消息授权时，可按需允许；关闭后仍可在消息中心主动查看。")

    add_heading(doc, "1.6 退款与异常处理", 1)
    add_data_table(
        doc,
        ["场景", "建议操作", "判定依据"],
        [
            ["支付后仍显示待支付", "返回订单列表下拉刷新，避免再次支付", "微信支付记录 + 订单最终状态"],
            ["退款中", "等待商家或平台审核，不重复提交", "售后单状态与消息通知"],
            ["二维码无法核销", "调高亮度，确认票券未使用且为对应场次", "核销端返回结果"],
            ["活动信息变化", "查看消息通知和活动详情，必要时联系主办方", "平台最新公告"],
        ],
        [1.45, 3.05, 2.0],
    )

    # Part 2: mini-program organizer
    add_part_title(doc, "PART 02", "小程序主办方端", "适用于主办方负责人、运营人员与现场核销员。")
    add_heading(doc, "2.1 入驻与进入管理后台", 1)
    add_step(doc, "在“我的”点击“我要入驻”，按页面要求填写主体与联系人信息并上传资质。")
    add_step(doc, "提交后在入驻页或消息中心查看审核进度；被驳回时先按原因修改，再重新提交。")
    add_step(doc, "审核通过后，从“我的 → 管理后台”进入主办方工作台。")
    add_paragraph(doc, "主办方底部导航通常包括“首页、活动、更多、账户”。核销员账号只应获得完成现场工作的必要权限。")
    add_figure_pair(
        doc,
        [
            ("mini_organizer_home.png", "主办方首页：活动与经营概览", "小程序主办方首页，展示活动和经营数据"),
            ("mini_organizer_activities.png", "活动中心：活动、销售、订单与核销", "小程序主办方活动中心，包含四个业务标签"),
        ],
    )

    add_heading(doc, "2.2 发布活动：五步流程", 1)
    add_paragraph(doc, "进入“活动 → 我的活动 → 新增活动”，依次完成活动信息、场地设定、上传海报、票券配置和活动资质。可保存草稿，确认完整后提交审核。")
    add_data_table(
        doc,
        ["步骤", "主要内容", "提交前检查"],
        [
            ["1 活动信息", "类型、名称、分享标题、日期、实名/未成年人规则、活动概要", "时间无冲突；限制说明清楚"],
            ["2 场地设定", "省市区、详细地址、经纬度或场地信息", "地图定位与文字地址一致"],
            ["3 上传海报", "详情页、列表页、分享等图片素材", "比例正确、文字清晰、无侵权"],
            ["4 票券配置", "票种、价格、库存、售卖时间、限购与退款条件", "金额、库存和时间逐项复核"],
            ["5 活动资质", "按活动类型上传备案或许可材料", "证件在有效期内且信息一致"],
        ],
        [1.1, 3.0, 2.4],
    )
    add_figure_pair(
        doc,
        [
            ("mini_organizer_create_step1.png", "活动信息填写", "小程序活动创建第一步，填写活动基本信息"),
            ("mini_organizer_create_step4.png", "票券配置", "小程序活动创建第四步，配置活动票券"),
        ],
    )
    add_callout(doc, "审核前最后检查", "活动一旦提交审核，部分字段可能暂时不可编辑。重点复核活动时间、地址、票价、库存、退款条件、实名和未成年人规则；宣传文案不得与票券配置矛盾。", kind="warn")

    add_heading(doc, "2.3 活动、销售和实时订单", 1)
    add_bullet(doc, "我的活动：按草稿、待审核、审核中、已上架、未通过等状态筛选；点击活动查看或编辑。")
    add_bullet(doc, "销售数据：按活动查看销量、订单金额等统计；统计口径以页面说明为准。")
    add_bullet(doc, "实时订单：按订单号、买家或状态搜索；涉及实名票时仅在业务需要范围内查看个人信息。")
    add_bullet(doc, "核销管理：维护核销员并查看核销记录；离场或离职人员应及时撤销权限。")

    add_heading(doc, "2.4 现场核销", 1)
    add_step(doc, "主办方先在核销管理中添加核销员；核销员按提示完成绑定。")
    add_step(doc, "活动开始前测试网络、相机权限与备用设备，并确认正在核销正确活动和场次。")
    add_step(doc, "扫描参与者票券二维码；成功后向参与者确认，失败时查看原因，不要连续快速重复扫码。")
    add_step(doc, "需要人工核对时，核实订单号、票种、使用状态和实名信息；异常票交由负责人处理。")
    add_callout(doc, "核销是不可逆业务动作", "只有参与者已到场且票券信息完全匹配时才确认核销。不要使用参与者转发的模糊截图，也不要在公开群聊收集完整二维码。", kind="danger")

    add_heading(doc, "2.5 账户与提现", 1)
    add_paragraph(doc, "在“账户”查看商家认证、总收益、可提现、冻结和已提现金额，并维护收款账户。")
    add_step(doc, "先提交或编辑收款账户，确保账户名、开户行与主体信息一致。")
    add_step(doc, "待收款账户审核通过、余额可提现后，再提交提现申请。")
    add_step(doc, "平台审核通过后进入打款流程；到账时间以页面显示的周期和银行处理为准。")
    add_bullet(doc, "等级规则由平台维护：LV1 服务费 5%，LV2 服务费 3%，LV3 服务费 0%。商家不可自行修改。")

    # Part 3: merchant web
    add_part_title(doc, "PART 03", "网页商家端", "适用于商家负责人、运营、客服、财务以及负责分配权限的管理员。")
    add_heading(doc, "3.1 登录与首次使用", 1)
    add_paragraph(doc, "使用平台提供的网页地址打开活动管理系统。默认进入商家登录页，可选择验证码登录或密码登录。")
    add_step(doc, "输入已入驻手机号，选择验证码或密码方式。验证码只能用于本人账号。")
    add_step(doc, "勾选商家服务协议后登录；如收不到验证码，使用页面上的帮助入口排查。")
    add_step(doc, "首次登录若要求设置密码，按规则完成设置；公共电脑使用完毕后务必退出。")
    add_figure(doc, "web_merchant_login.png", "网页商家端登录页", width=6.2,
               alt="Hyper 网页商家端登录页，支持验证码和密码登录")

    add_heading(doc, "3.2 首页与经营概览", 1)
    add_paragraph(doc, "首页左侧显示活动，右侧展示总活动数、订单金额、今日订单金额、销量等指标，并提供常用业务入口。")
    add_bullet(doc, "点击已上架活动可查看活动或数据；待审核活动以状态提示为准。")
    add_bullet(doc, "快速配置可进入新增活动、添加核销员、活动合集、收款账户和消息中心。")
    add_bullet(doc, "指标受筛选、支付、退款与数据刷新时间影响；对账时以订单和资金明细为准。")
    add_figure(doc, "web_merchant_dashboard.png", "商家首页：活动与数据概览", width=6.2,
               alt="网页商家首页，左侧是活动，右侧是经营数据和快捷入口")

    add_heading(doc, "3.3 活动管理与发布", 1)
    add_paragraph(doc, "左侧点击“活动”。顶部标签可在“我的活动、销售数据、实时订单、核销管理”之间切换。")
    add_step(doc, "在“我的活动”使用关键词、主办方、发布日期、活动日期与状态进行筛选。")
    add_step(doc, "点击“新增活动”，按照五步向导录入信息；右侧手机预览用于校对小程序展示。")
    add_step(doc, "保存草稿后可继续编辑；提交审核前按五步检查清单逐项复核。")
    add_step(doc, "未通过时打开详情查看驳回原因，修改对应字段或材料后重新提交。")
    add_figure(doc, "web_merchant_events.png", "活动列表与筛选", width=6.2,
               alt="网页商家端活动列表，可按状态和日期筛选")
    add_figure(doc, "web_merchant_event_wizard.png", "五步活动发布向导与手机预览", width=6.2,
               alt="网页活动发布向导，包含五个步骤和右侧手机预览")

    add_heading(doc, "3.4 订单与售后", 1)
    add_paragraph(doc, "点击“订单售后”，可按关键词、订单状态、活动、退款状态和下单日期筛选。点击订单行后，右侧显示订单与售后详情。")
    add_step(doc, "处理前核对订单号、活动、票种、数量、实付金额、买家和使用状态。")
    add_step(doc, "进入“售后订单”查看退款原因与金额；根据活动规则和实际情况选择通过或驳回。")
    add_step(doc, "驳回时填写清楚原因；通过后以退款状态为准，避免重复操作。")
    add_step(doc, "待支付订单的取消、已核销票券的售后等边界场景，应按页面允许的动作处理，不绕过系统。")
    add_figure(doc, "web_merchant_orders.png", "商家订单与售后工作台", width=6.2,
               alt="网页商家订单与售后页面，展示订单筛选、列表和详情")
    add_callout(doc, "隐私最小化", "订单页可能显示脱敏手机号与实名信息。仅为履约、核销和售后目的查看，不得下载、复制或传播到无关群聊。", kind="warn")

    add_heading(doc, "3.5 账户、收款与提现", 1)
    add_paragraph(doc, "点击“账户资金”查看商家等级、服务费、总收益、可提现、冻结金额、已提现以及收款账户状态。")
    add_step(doc, "首次收款先编辑收款账户并提交审核；审核通过前不可发起提现。")
    add_step(doc, "点击“申请提现”，核对可提现余额与收款账户后提交。")
    add_step(doc, "页面显示 T+1 表示平台审核通过后进入打款队列；实际到账仍受银行处理影响。")
    add_figure(doc, "web_merchant_account.png", "账户与资金：等级、余额和收款账户", width=6.2,
               alt="网页商家账户资金页面，包含等级规则、余额和收款账户")
    add_callout(doc, "等级规则只读", "等级与服务费由平台统一维护。商家可查看当前等级及权益，但不能编辑规则；如对等级或费用有疑问，通过消息中心或平台支持渠道反馈。", kind="info")

    add_heading(doc, "3.6 运营中心与活动合集", 1)
    add_bullet(doc, "消息中心：查看审核、费用和系统通知；点击消息或“一键已读”更新未读状态。")
    add_bullet(doc, "动态管理：发布与活动相关的动态，或隐藏不再展示的内容；发布前确认素材授权。")
    add_bullet(doc, "活动合集：把多个活动整理为专题集合，用于统一展示和运营。")
    add_figure(doc, "web_merchant_ops.png", "运营中心：消息与动态管理", width=6.2,
               alt="网页商家运营中心，包含消息中心和动态管理")

    add_heading(doc, "3.7 权限设置", 1)
    add_paragraph(doc, "“权限设置”包含角色管理、子账号和操作日志。负责人应先定义角色，再为子账号分配最小必要权限。")
    add_step(doc, "在角色管理中新建角色，填写名称、说明并勾选活动、订单售后、核销、财务、消息等权限。")
    add_step(doc, "在子账号页新增成员并绑定角色；岗位变化或离职后立即停用或调整。")
    add_step(doc, "通过操作日志核查关键动作；如发现异常，先暂停相关账号并保留日志。")
    add_figure(doc, "web_merchant_access.png", "角色管理与权限勾选", width=6.2,
               alt="网页商家权限设置页，展示角色列表与新增角色表单")
    add_callout(doc, "建议分工", "运营角色管理活动与消息；客服角色处理订单售后；核销员只保留核销权限；财务角色查看资金和提现；只有负责人保留角色与子账号管理权限。", kind="success")

    # Part 4: admin web
    add_part_title(doc, "PART 04", "平台管理端", "仅适用于获得平台后台权限的管理员；菜单会按账号权限动态显示。")
    add_heading(doc, "4.1 管理员登录与权限", 1)
    add_paragraph(doc, "在商家登录页点击“管理员登录”，输入管理员账号和密码。登录后，只会看到当前账号有权访问的菜单。")
    add_step(doc, "确认进入的是“管理员登录”页面，不要在商家验证码登录框输入管理员凭据。")
    add_step(doc, "登录后核对左侧菜单；缺少业务菜单时联系具备权限管理能力的负责人，不要借用他人账号。")
    add_step(doc, "处理完敏感业务后退出登录；公共设备不保存密码。")
    add_figure(doc, "web_admin_login.png", "平台管理员登录页", width=6.2,
               alt="Hyper 平台管理员账号密码登录页")

    add_heading(doc, "4.2 数据概览", 1)
    add_paragraph(doc, "概览页集中显示派对/场地、活动、票券、订单、用户、总收入和商家数，用于快速判断平台规模与异常波动。")
    add_bullet(doc, "概览数字用于运营观察；需要核对具体交易时进入订单、售后或财务模块。")
    add_bullet(doc, "数据权限与菜单权限相互配合，敏感数据只在工作需要范围内查看。")
    add_figure(doc, "web_admin_dashboard.png", "平台数据概览与权限菜单", width=6.2,
               alt="平台管理端概览页，展示业务总量和左侧权限菜单")

    add_heading(doc, "4.3 入驻与活动审核", 1)
    add_paragraph(doc, "管理员可在具备权限时处理入驻审核和活动审核。审核时应以页面材料、平台规则和必要的外部核验为依据。")
    add_step(doc, "在审核列表按名称、ID、日期或状态搜索目标记录。")
    add_step(doc, "打开详情，核对主体、联系人、资质有效期，以及活动时间、地址、海报、票券和限制说明。")
    add_step(doc, "通过时确认信息完整；驳回时填写具体、可执行的修改原因。")
    add_step(doc, "提交后回到列表确认状态已更新，并关注是否产生消息通知。")
    add_figure(doc, "web_admin_activity_audits.png", "活动审核列表与状态筛选", width=6.2,
               alt="平台活动审核列表，可按待审核、审核中、已上架和未通过筛选")
    add_callout(doc, "审核一致性", "相同类型活动应采用一致标准。驳回原因应明确指出具体字段或材料，避免使用笼统表述；涉及高风险资质时，应按照业务审批规范升级复核。", kind="warn")

    add_heading(doc, "4.4 订单、售后与核销", 1)
    add_paragraph(doc, "订单页提供关键词、状态和活动 ID 等筛选；同页可查看售后订单。核销页用于查询核销结果与记录。")
    add_step(doc, "处理订单问题前先锁定唯一订单号，核对用户、活动、金额、状态和创建时间。")
    add_step(doc, "售后审核先查看商家处理信息、退款原因与票券使用状态，再选择通过或驳回。")
    add_step(doc, "核销争议以系统记录、票券状态、扫码时间和操作账号为依据，必要时保留证据。")
    add_figure(doc, "web_admin_orders.png", "平台订单列表与售后审核", width=6.2,
               alt="平台管理端订单列表和售后审核区域")

    doc.add_page_break()
    add_heading(doc, "4.5 其他管理模块", 1)
    add_data_table(
        doc,
        ["模块", "主要用途", "操作要点"],
        [
            ["系统", "平台配置与基础规则", "变更前确认影响范围，保留配置依据"],
            ["用户 / 商家 / 主办方资料", "查询与维护主体信息", "仅处理有依据的变更，保护个人信息"],
            ["票券", "查询票种、库存与使用状态", "金额、库存和使用状态需要交叉核对"],
            ["活动合集", "平台专题集合管理", "上线前校验活动有效性与排序"],
            ["动态消息", "内容与通知治理", "按规则处理不合规内容，保留处理记录"],
            ["财务积分", "提现、资金和积分相关业务", "线下打款完成后再确认“已打款”"],
        ],
        [1.35, 2.25, 2.9],
        trailing_space=False,
    )
    add_callout(
        doc,
        "财务操作红线",
        "提现审核通过不等于银行已完成转账。只有在真实线下打款完成并核对金额与收款账户后，才在系统确认“已打款”；关键操作建议双人复核。",
        kind="danger",
        compact=True,
        trailing_space=False,
    )

    # Part 5: appendix
    add_part_title(doc, "PART 05", "状态与常见问题", "用于快速判断当前进度、排查异常并降低误操作风险。")
    add_heading(doc, "5.1 常用状态速查", 1)
    add_data_table(
        doc,
        ["对象", "状态", "含义与下一步"],
        [
            ["活动", "草稿", "活动信息尚处于编辑状态；请完善各项内容后提交审核"],
            ["活动", "待审核 / 审核中", "已进入平台流程；等待结果，避免重复提交"],
            ["活动", "已上架", "用户可见；持续关注售卖、通知和现场准备"],
            ["活动", "未通过", "打开详情查看原因，修改后重新提交"],
            ["订单", "待支付", "尚未完成支付；在有效期内支付或等待关闭"],
            ["订单", "待使用", "已支付且未核销；到场出示电子票"],
            ["订单", "已使用", "已完成核销，票券不可重复使用"],
            ["售后", "待审核 / 退款中", "等待商家、平台或支付渠道处理"],
            ["售后", "已退款 / 已驳回", "流程结束；查看金额、原因与通知"],
            ["收款账户", "审核中 / 已通过", "审核中不可提现；通过后按余额申请"],
        ],
        [1.0, 1.7, 3.8],
    )

    add_heading(doc, "5.2 常见问题", 1)
    qa_items = [
        ("小程序地图无法显示或定位失败如何处理？", "请确认微信定位权限、系统定位服务和网络连接。拒绝定位授权时，仍可使用列表、搜索和筛选功能；如问题持续，请退出后重新进入或联系平台支持。"),
        ("支付成功后订单状态未更新如何处理？", "请返回订单列表刷新，并核对微信支付记录，避免重复支付。如状态仍未更新，请记录订单号和支付时间，并通过平台支持渠道反馈。"),
        ("活动提交后无法编辑如何处理？", "活动处于待审核或审核中状态时，部分字段可能被锁定。请等待审核结果；如需紧急修改，请按照平台流程联系管理员。"),
        ("网页端未显示所需菜单如何处理？", "菜单由角色权限控制。请联系商家负责人或平台权限管理员调整权限，不得共用他人账号。"),
        ("退款长时间未到账如何处理？", "请先核对售后状态。商家或平台审核通过后，支付渠道仍可能需要一定处理时间；最终结果以退款记录和原支付渠道到账信息为准。"),
        ("核销扫码失败如何处理？", "请确认相机权限、网络连接、活动场次和票券状态，并请参与者打开原始电子票、调高屏幕亮度，避免连续重复扫码。"),
        ("商家等级规则为何不可编辑？", "等级和服务费由平台统一维护，商家端仅提供查看功能。当前规则为 LV1 服务费 5%、LV2 服务费 3%、LV3 服务费 0%。"),
        ("是否支持活动分销或活动抽奖？", "当前版本暂不提供活动分销和活动抽奖功能。如需集中展示多个活动，可使用活动合集功能。"),
    ]
    for question, answer in qa_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(2)
        set_keep_with_next(p)
        r = p.add_run(f"问题  {question}")
        set_run_font(r, size=10.5, bold=True, color=BLUE_DARK)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.22)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(f"处理建议  {answer}")
        set_run_font(r2, size=10, color=INK)

    add_heading(doc, "5.3 操作与安全清单", 1)
    add_data_table(
        doc,
        ["时点", "检查事项"],
        [
            ["发布活动前", "时间、地点、票价、库存、实名/未成年规则、退款条件、海报与资质一致"],
            ["处理退款前", "订单号、实付金额、使用状态、退款原因和适用规则一致"],
            ["现场核销前", "账号权限、正确活动/场次、网络、相机、备用设备可用"],
            ["提交提现前", "可提现余额、账户名、银行账号和审核状态正确"],
            ["确认已打款前", "真实转账已完成，金额与账户核对无误，并完成内部复核"],
            ["人员变动时", "立即停用子账号或核销员权限，检查操作日志"],
        ],
        [1.55, 4.95],
    )

    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(130)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("支持与反馈")
    set_run_font(r, size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("如需反馈问题，请记录相关页面、发生时间、账号角色、订单或活动编号及当前状态，并通过消息中心或指定支持渠道提交。")
    set_run_font(r, size=11.5, color=BLUE_DARK)
    accent = doc.add_table(rows=1, cols=3)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent.autofit = False
    remove_table_borders(accent)
    for idx, fill in enumerate((GOLD, PINK, PALE_BLUE)):
        c = accent.cell(0, idx)
        c.width = Inches(2.15)
        set_cell_shading(c, fill)
        set_cell_margins(c, top=28, start=0, bottom=28, end=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("Hyper 用户使用手册  ·  2026年8月版")
    set_run_font(r, size=9.5, bold=True, color=MUTED)
    return doc


if __name__ == "__main__":
    manual = build_manual()
    manual.save(OUTPUT_PATH)
    print(OUTPUT_PATH)
